from dotenv import load_dotenv
load_dotenv()

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

"""
NIIFL Demo Runner
==================
Single command that starts the full demo environment.

Usage:
    python3 demo/demo_runner.py

What it does:
    1. Validates API key and skills are ready
    2. Seeds pre-polished demo deals into the server
    3. Starts FastAPI web server on port 8000
    4. Opens browser automatically

On demo day: run this ONE command, then open http://localhost:8000
"""

import os, sys, json, asyncio, datetime, webbrowser, time
from pathlib import Path

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "config"))
sys.path.insert(0, str(ROOT / "mcp_servers"))
sys.path.insert(0, str(ROOT / "demo"))

OUTPUTS_DIR = ROOT / "outputs"
OUTPUTS_DIR.mkdir(exist_ok=True)
COMPLETED_DEALS_FILE = OUTPUTS_DIR / "completed_deals.json"

_DEAL_DOCX_NAMES = {
    "DC001": "GHL_IC_Memo.docx",
    "DC002": "SunGrid_IC_Memo.docx",
    "DC003": "IndiaStack_IC_Memo.docx",
    "DC004": "Sunrise_IC_Memo.docx",
}

def _build_docx(deal_id: str, deal: dict, out_path: str) -> str:
    """
    Generate a formatted IC memo DOCX using python-docx.
    Handles all Unicode characters correctly on Windows.
    Returns out_path.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    FUND_DISPLAY = {
        "master_fund":                  "NIIF Master Fund",
        "strategic_opportunities_fund": "NIIF Strategic Opportunities Fund",
        "private_markets_fund":         "NIIF Private Markets Fund",
        "india_japan_fund":             "NIIF India-Japan Fund",
    }
    NAVY   = RGBColor(0x1F, 0x38, 0x64)
    BLUE   = RGBColor(0x2E, 0x74, 0xB5)
    RED    = RGBColor(0xC0, 0x00, 0x00)
    GREY   = RGBColor(0x59, 0x59, 0x59)

    memo   = deal.get("memo_sections") or {}
    grader = deal.get("grader") or {}
    score  = grader.get("quality_score", 0)
    fund   = FUND_DISPLAY.get(deal.get("fund",""), deal.get("fund",""))
    name   = deal.get("name", "")
    today  = datetime.date.today().strftime("%d %B %Y")

    SECTIONS = [
        ("1. Investment Thesis",         "thesis"),
        ("2. Market & Competitive Position", "market_analysis"),
        ("3. Financial Returns",          "financial_returns"),
        ("4. Key Risks & Mitigants",      "risks_mitigants"),
        ("5. ESG & Impact",               "esg_impact"),
        ("6. Policy Alignment",           "policy_alignment"),
        ("7. Deal Structure & Terms",     "deal_structure"),
        ("8. Recommendation",             "recommendation"),
    ]

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def set_run_font(run, size=11, bold=False, color=None, italic=False):
        run.font.name  = "Arial"
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        set_run_font(run, size=14 if level == 1 else 12, bold=True,
                     color=NAVY if level == 1 else BLUE)
        return p

    def add_body(text):
        if not text:
            text = "[To be completed]"
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(str(text))
        set_run_font(run, size=10)
        return p

    def add_kv(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=10, bold=True)
        r2 = p.add_run(str(value) if value is not None else "—")
        set_run_font(r2, size=10)
        return p

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E74B5")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ── Cover ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("CONFIDENTIAL")
    set_run_font(r, size=10, bold=True, color=RED)

    p = doc.add_paragraph()
    r = p.add_run("Investment Committee Memorandum")
    set_run_font(r, size=22, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(6)

    add_divider()
    add_kv("Company",         name)
    add_kv("Fund",            fund)
    add_kv("Deal ID",         deal_id)
    add_kv("Date",            today)
    add_kv("AI Quality Score", f"{score}/100")
    add_kv("Base IRR",        f"{deal.get('irr_base', '—')}%")
    add_kv("Base MOIC",       f"{deal.get('moic', '—')}x")
    doc.add_page_break()

    # ── Memo Sections ──────────────────────────────────────────────────────────
    for heading, key in SECTIONS:
        add_heading(heading, level=1)
        add_body(memo.get(key, ""))
        add_divider()

    # ── Returns table ──────────────────────────────────────────────────────────
    add_heading("Returns Summary", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Scenario", "Gross IRR", "MOIC"]
    rows_data = [
        ("Base case", f"{deal.get('irr_base','—')}%", f"{deal.get('moic','—')}x"),
        ("Bull case", f"{deal.get('irr_bull','—')}%",  "—"),
        ("Bear case", f"{deal.get('irr_bear','—')}%",  "—"),
    ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
    for r_idx, (s, irr, moic) in enumerate(rows_data, start=1):
        for c_idx, val in enumerate([s, irr, moic]):
            table.rows[r_idx].cells[c_idx].text = val
    doc.add_paragraph()

    # ── QA Appendix ────────────────────────────────────────────────────────────
    add_heading("Appendix — AI Quality Assurance", level=1)
    add_kv("Overall Pass",   grader.get("overall_pass", False))
    add_kv("Quality Score",  f"{score}/100")
    add_kv("Grader Notes",   grader.get("grader_notes", ""))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(
        "— This memo was drafted by NIIFL's AI engine and reviewed by a qualified "
        "investment professional before IC presentation. —"
    )
    set_run_font(r, size=8, italic=True, color=GREY)

    doc.save(out_path)
    return out_path


def _generate_startup_docx(deal_id: str, deal: dict) -> str | None:
    """Generate a DOCX for a pre-loaded deal at startup. Returns path or None on failure."""
    filename = _DEAL_DOCX_NAMES.get(deal_id, f"{deal_id}_IC_Memo.docx")
    out_path = str(OUTPUTS_DIR / filename)
    if Path(out_path).exists():
        return out_path  # already generated this session
    try:
        _build_docx(deal_id, deal, out_path)
        return out_path
    except Exception as e:
        print(f"  [DOCX] Warning — could not generate {filename}: {e}")
        return None

# ── Completed deal persistence ────────────────────────────────────────────────
def _save_completed_deal(deal_id: str, deal: dict):
    """Append/update a completed pipeline deal in completed_deals.json."""
    try:
        existing = {}
        if COMPLETED_DEALS_FILE.exists():
            with open(COMPLETED_DEALS_FILE, "r", encoding="utf-8") as f:
                existing = json.load(f)
        existing[deal_id] = {k: v for k, v in deal.items()
                             if isinstance(v, (str, int, float, bool, dict, list, type(None)))}
        with open(COMPLETED_DEALS_FILE, "w", encoding="utf-8") as f:
            json.dump(existing, f, indent=2, default=str, ensure_ascii=False)
        print(f"[Pipeline] Saved {deal_id} to completed_deals.json")
    except Exception as e:
        print(f"[Pipeline] Could not save deal to completed_deals.json: {e}")

def _load_completed_deals() -> dict:
    """Load previously completed pipeline deals from completed_deals.json."""
    if not COMPLETED_DEALS_FILE.exists():
        return {}
    try:
        with open(COMPLETED_DEALS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"  [Startup] Could not load completed_deals.json: {e}")
        return {}

# ── Startup checks ─────────────────────────────────────────────────────────────
def check_ready():
    print("\n" + "="*52)
    print("  NIIFL AI Engine — Demo Mode")
    print("="*52)

    api_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not api_key or api_key == "YOUR_API_KEY_HERE":
        print("\n  ERROR: ANTHROPIC_API_KEY not set.")
        print("  Run: export ANTHROPIC_API_KEY=sk-ant-...")
        sys.exit(1)
    print(f"  API key:  {api_key[:14]}...{api_key[-4:]}  OK")

    skills_ok = (ROOT / "skills" / "skill_fund_mandate.txt").exists()
    if not skills_ok:
        print("  Skills missing — building now...")
        os.system(f"cd {ROOT} && python3 step1_build_skills.py")
    print(f"  Skills:   {'OK' if skills_ok else 'built'}")

    mcp_ok = (ROOT / "mcp_servers" / "mcp_dealcloud.py").exists()
    if not mcp_ok:
        print("  MCP servers missing — building now...")
        os.system(f"cd {ROOT} && python3 step2_build_mcp_servers.py")
    print(f"  MCP:      {'OK' if mcp_ok else 'built'}")

    print("\n  Pre-loaded deals: Greenway Highways, SunGrid Renewables, IndiaStack")
    print("  Live demo deal:   Sunrise Port Terminal (submit during demo)")
    print(f"\n  Starting server → http://localhost:8000")
    print("="*52 + "\n")

# ── Patch web/main.py to use demo data ────────────────────────────────────────
def build_demo_server():
    """
    Imports the FastAPI app from web/main.py and replaces
    DEALS with our pre-polished demo data.
    """
    from demo_data import DEMO_DEALS, LIVE_DEMO_BRIEF
    import web.main as web_main

    # Replace DEALS with demo data
    web_main.DEALS.clear()
    web_main.DEALS.update(DEMO_DEALS)

    # Restore completed pipeline deals that survived previous sessions
    completed = _load_completed_deals()
    print(f"  Loaded {len(completed)} completed deal(s) from storage")
    for did, deal in completed.items():
        # Verify DOCX still on disk; clear flag if file was deleted
        if deal.get("docx_path") and not Path(deal["docx_path"]).exists():
            deal["docx_ready"] = False
        web_main.DEALS[did] = deal
        print(f"    {did} ({deal.get('name','?')}) restored — score {deal.get('quality_score','?')}/100")

    # Patch the pipeline to use the demo agent
    original_run = web_main.run_pipeline_background

    async def demo_pipeline(deal_id: str, submission):
        """Uses the streamlined demo agent instead of the full pipeline."""
        deal = web_main.DEALS[deal_id]

        _HURDLE = {
            "master_fund":                  12.0,
            "india_japan_fund":             14.0,
            "strategic_opportunities_fund": 15.0,
            "private_markets_fund":         12.0,
        }
        hurdle_irr = _HURDLE.get(submission.fund, 12.0)

        # Update agent status for UI — reset all, then mark orchestrator running
        for agent in web_main.AGENT_STATUS:
            web_main.AGENT_STATUS[agent]["status"] = "idle"
            web_main.AGENT_STATUS[agent]["completed_at"] = None
        web_main.AGENT_STATUS["orchestrator"]["status"] = "running"
        web_main.AGENT_STATUS["orchestrator"]["current_deal"] = deal["name"]

        deal["status"] = "screening"
        web_main.AGENT_STATUS["deal_sourcing"]["status"] = "running"
        web_main.AGENT_STATUS["deal_sourcing"]["current_deal"] = deal["name"]
        web_main.AGENT_STATUS["deal_sourcing"]["completed_at"] = None

        async def update_progress(stage: str, pct: int, msg: str = ""):
            stage_map = {
                "screening": "deal_sourcing",
                "modelling": "financial_modelling",
                "memo":      "ic_memo_drafting",
                "grading":   "outcomes_grader",
            }
            agent_key = stage_map.get(stage, stage)

            # Ensure pipeline_stages exists on the deal
            if "pipeline_stages" not in deal:
                deal["pipeline_stages"] = {}
            ps = deal["pipeline_stages"]
            if stage not in ps:
                ps[stage] = {"status": "pending", "started_at": None, "completed_at": None, "duration_s": None, "msg": ""}

            ps[stage]["msg"] = msg

            if pct < 100:
                # Stage starting or in progress
                if ps[stage]["status"] != "running":
                    ps[stage]["status"] = "running"
                    ps[stage]["started_at"] = time.time()
                web_main.AGENT_STATUS[agent_key]["status"] = "running"
                web_main.AGENT_STATUS[agent_key]["current_deal"] = deal["name"]
                web_main.AGENT_STATUS[agent_key]["completed_at"] = None
                stage_to_deal_status = {
                    "screening": "screening",
                    "modelling": "dd",
                    "memo": "ic_prep",
                    "grading": "ic_prep",
                }
                deal["status"] = stage_to_deal_status.get(stage, deal["status"])
            else:
                # Stage complete
                _now = time.time()
                ps[stage]["status"] = "done"
                ps[stage]["completed_at"] = _now
                if ps[stage].get("started_at"):
                    ps[stage]["duration_s"] = round(_now - ps[stage]["started_at"], 1)
                web_main.AGENT_STATUS[agent_key]["status"] = "idle"
                web_main.AGENT_STATUS[agent_key]["current_deal"] = None
                web_main.AGENT_STATUS[agent_key]["completed_at"] = _now
                # Append to activity feed
                _ACT = {
                    "screening": ("ti-search",    "var(--blue-bg)",   "var(--blue)",   "Deal sourcing agent"),
                    "modelling": ("ti-chart-bar", "#dbeafe",          "#1e40af",       "Financial modelling agent"),
                    "memo":      ("ti-file-text", "var(--purple-bg)", "var(--purple)", "IC memo agent"),
                    "grading":   ("ti-check",     "var(--green-bg)",  "var(--green)",  "Outcomes grader"),
                }
                if stage in _ACT:
                    icon, bg, color, agent_label = _ACT[stage]
                    if stage == "screening":
                        act_text = f"Deal screened — {deal['name']} passed initial screening"
                    elif stage == "modelling":
                        irr_str = msg.split("Base IRR ")[-1] if "Base IRR" in msg else "?"
                        act_text = f"Model complete — {deal['name']} base IRR {irr_str} vs {hurdle_irr}% hurdle"
                    elif stage == "memo":
                        act_text = f"Memo drafted — {deal['name']} 8-section memo complete"
                    else:
                        score_str = msg.split("score ")[-1] if "score" in msg else "?"
                        act_text = f"IC memo ready — {deal['name']} passed QA ({score_str}). Partner review required."
                    web_main.ACTIVITY_LOG.append({
                        "type": stage, "text": act_text,
                        "icon": icon, "bg": bg, "color": color,
                        "agent": agent_label, "ts": _now,
                    })

        from demo_agent import run_demo_pipeline
        financials = {
            "revenue_cr": {
                "FY22": submission.revenue_fy22,
                "FY23": submission.revenue_fy23,
                "FY24": submission.revenue_fy24,
            },
            "ebitda_cr": {
                "FY22": submission.ebitda_fy22,
                "FY23": submission.ebitda_fy23,
                "FY24": submission.ebitda_fy24,
            },
        }

        result = await run_demo_pipeline(
            deal_name=submission.name,
            fund=submission.fund,
            sector=submission.sector,
            ticket_cr=submission.ticket_cr,
            stake_pct=submission.stake_pct,
            entry_ev_cr=submission.entry_ev_cr,
            brief=submission.brief,
            financials=financials,
            progress_callback=update_progress,
        )

        # Update deal with results
        deal["memo_sections"]  = result.get("memo_sections", {})
        deal["grader"]         = result.get("grader")
        deal["quality_score"]  = result.get("quality_score")
        deal["irr_base"]       = result.get("irr_base")
        deal["irr_bull"]       = result.get("irr_bull")
        deal["irr_bear"]       = result.get("irr_bear")
        deal["irr_net"]        = result.get("irr_net")
        deal["moic"]           = result.get("moic")
        deal["memo_ready"] = bool(deal.get("memo_sections"))
        deal["status"]     = "partner_review" if (result.get("grader") or {}).get("overall_pass") else "ic_prep"

        # Always generate DOCX using python-docx (handles Unicode on Windows)
        try:
            filename = _DEAL_DOCX_NAMES.get(deal_id, f"{deal_id}_IC_Memo.docx")
            out_path = str(OUTPUTS_DIR / filename)
            _build_docx(deal_id, deal, out_path)
            deal["docx_path"] = out_path
            deal["docx_ready"] = True
            print(f"[Pipeline] DOCX written: {filename}")
        except Exception as e:
            deal["docx_error"] = str(e)
            print(f"[Pipeline] DOCX generation failed: {e}")

        # Persist completed deal so it survives server restarts
        _save_completed_deal(deal_id, deal)

        # Reset orchestrator
        import time as _t
        web_main.AGENT_STATUS["orchestrator"]["status"] = "idle"
        web_main.AGENT_STATUS["orchestrator"]["current_deal"] = None
        web_main.AGENT_STATUS["orchestrator"]["completed_at"] = _t.time()

    # Monkey-patch the background task
    web_main.run_pipeline_background = demo_pipeline

    # Ensure DOCXs exist for all pre-loaded deals. Skips files already on disk.
    print("  Checking DOCX files for pre-loaded deals...")
    for did, deal in web_main.DEALS.items():
        if deal.get("memo_ready") and deal.get("memo_sections"):
            filename = _DEAL_DOCX_NAMES.get(did, f"{did}_IC_Memo.docx")
            out_path = str(OUTPUTS_DIR / filename)
            if Path(out_path).exists():
                deal["docx_path"] = out_path
                deal["docx_ready"] = True
                print(f"    {did} ({deal['name']}) already exists, skipping")
            else:
                path = _generate_startup_docx(did, deal)
                if path:
                    deal["docx_path"] = path
                    deal["docx_ready"] = True
                    print(f"    {did} ({deal['name']}) generated: {filename}")
                else:
                    print(f"    {did} ({deal['name']}) DOCX generation failed")

    # Reset any deals left in a mid-pipeline status from a previous session.
    # Statuses that mean "pipeline was running when server died":
    _RUNNING_STATUSES = {"queued", "screening", "dd", "ic_prep"}
    for deal in web_main.DEALS.values():
        if deal.get("status") in _RUNNING_STATUSES and not deal.get("memo_ready"):
            deal["status"] = "ic_prep"
            deal["memo_ready"] = True
            if not deal.get("memo_sections"):
                deal["memo_sections"] = {
                    "thesis": f"Pipeline for {deal.get('name','this deal')} was interrupted by a server restart. Re-submit to generate a full IC memo.",
                }
            if not deal.get("grader"):
                deal["grader"] = {"overall_pass": False, "quality_score": 0,
                                   "grader_notes": "Pipeline interrupted — resubmit to complete."}
            print(f"  [Startup] Reset stuck deal '{deal.get('name')}' → ic_prep")

    return web_main.app

# ── Main ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    check_ready()
    app = build_demo_server()
    port = int(os.getenv("PORT", 8000))

    # Open browser after short delay
    def open_browser():
        time.sleep(2.5)
        webbrowser.open(f"http://localhost:{port}")
        print(f"\n  Browser opened → http://localhost:{port}")
        print("  Demo is ready. Pipeline view shows 3 pre-loaded deals.")
        print("  Submit 'Sunrise Port Terminal' live when ready.\n")

    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="warning")
