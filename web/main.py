"""
NIIFL AI Engine — FastAPI Backend (Demo Mode)
=============================================
Entry point for Render deployment.

Render start command: uvicorn web.main:app --host 0.0.0.0 --port $PORT

Pipeline uses demo_agent.py (parallel Sonnet section drafting).
Demo deals are seeded from demo_data.py at startup.
Completed pipeline deals are persisted to outputs/completed_deals.json.
"""

import os
import sys
import json
import time
import asyncio
import datetime
from pathlib import Path
from typing import Optional

# ── Path setup ────────────────────────────────────────────────────────────────
WEB_DIR  = Path(__file__).parent
ROOT_DIR = WEB_DIR.parent
sys.path.insert(0, str(ROOT_DIR / "demo"))   # demo_agent.py, demo_data.py, demo_runner.py
sys.path.insert(0, str(ROOT_DIR))
sys.path.insert(0, str(ROOT_DIR / "config"))
sys.path.insert(0, str(ROOT_DIR / "mcp_servers"))

from fastapi import FastAPI, BackgroundTasks, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# ── Demo data ─────────────────────────────────────────────────────────────────
from demo_data import DEMO_DEALS

# ── Persistence paths ─────────────────────────────────────────────────────────
OUTPUTS_DIR          = ROOT_DIR / "outputs"
OUTPUTS_DIR.mkdir(exist_ok=True)
COMPLETED_DEALS_FILE = OUTPUTS_DIR / "completed_deals.json"

_DEAL_DOCX_NAMES = {
    "DC001": "GHL_IC_Memo.docx",
    "DC002": "SunGrid_IC_Memo.docx",
    "DC003": "IndiaStack_IC_Memo.docx",
    "DC004": "Sunrise_IC_Memo.docx",
}

_HURDLE = {
    "master_fund":                  12.0,
    "india_japan_fund":             14.0,
    "strategic_opportunities_fund": 15.0,
    "private_markets_fund":         12.0,
}

# ── DOCX builder ──────────────────────────────────────────────────────────────
def _build_docx(deal_id: str, deal: dict, out_path: str) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FUND_DISPLAY = {
        "master_fund":                  "Master Fund",
        "strategic_opportunities_fund": "Strategic Opportunities Fund",
        "private_markets_fund":         "Private Markets Fund",
        "india_japan_fund":             "India-Japan Fund",
    }
    NAVY = RGBColor(0x1F, 0x38, 0x64)
    BLUE = RGBColor(0x2E, 0x74, 0xB5)
    RED  = RGBColor(0xC0, 0x00, 0x00)
    GREY = RGBColor(0x59, 0x59, 0x59)

    memo   = deal.get("memo_sections") or {}
    grader = deal.get("grader") or {}
    score  = grader.get("quality_score", 0)
    fund   = FUND_DISPLAY.get(deal.get("fund", ""), deal.get("fund", ""))
    name   = deal.get("name", "")
    today  = datetime.date.today().strftime("%d %B %Y")

    SECTIONS = [
        ("1. Investment Thesis",             "thesis"),
        ("2. Market & Competitive Position", "market_analysis"),
        ("3. Financial Returns",             "financial_returns"),
        ("4. Key Risks & Mitigants",         "risks_mitigants"),
        ("5. ESG & Impact",                  "esg_impact"),
        ("6. Policy Alignment",              "policy_alignment"),
        ("7. Deal Structure & Terms",        "deal_structure"),
        ("8. Recommendation",               "recommendation"),
    ]

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.5)

    def _font(run, size=11, bold=False, color=None, italic=False):
        run.font.name   = "Arial"
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

    def add_heading(text, level=1):
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        _font(run, size=14 if level == 1 else 12, bold=True,
              color=NAVY if level == 1 else BLUE)
        return p

    def add_body(text):
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(str(text) if text else "[To be completed]")
        _font(run, size=10)
        return p

    def add_kv(label, value):
        p  = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        _font(r1, size=10, bold=True)
        r2 = p.add_run(str(value) if value is not None else "—")
        _font(r2, size=10)
        return p

    def add_divider():
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "2E74B5")
        pBdr.append(bot)
        pPr.append(pBdr)
        return p

    p = doc.add_paragraph()
    _font(p.add_run("CONFIDENTIAL"), size=10, bold=True, color=RED)

    p = doc.add_paragraph()
    _font(p.add_run("Investment Committee Memorandum"), size=22, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(6)

    add_divider()
    add_kv("Company",          name)
    add_kv("Fund",             fund)
    add_kv("Deal ID",          deal_id)
    add_kv("Date",             today)
    add_kv("AI Quality Score", f"{score}/100")
    add_kv("Base IRR",         f"{deal.get('irr_base', '—')}%")
    add_kv("Base MOIC",        f"{deal.get('moic', '—')}x")
    doc.add_page_break()

    for heading, key in SECTIONS:
        add_heading(heading, level=1)
        add_body(memo.get(key, ""))
        add_divider()

    add_heading("Returns Summary", level=2)
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Gross IRR", "MOIC"]):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
    for r_idx, (s, irr, moic) in enumerate([
        ("Base case", f"{deal.get('irr_base','—')}%", f"{deal.get('moic','—')}x"),
        ("Bull case", f"{deal.get('irr_bull','—')}%", "—"),
        ("Bear case", f"{deal.get('irr_bear','—')}%", "—"),
    ], start=1):
        for c_idx, val in enumerate([s, irr, moic]):
            tbl.rows[r_idx].cells[c_idx].text = val
    doc.add_paragraph()

    add_heading("Appendix — AI Quality Assurance", level=1)
    add_kv("Overall Pass",  grader.get("overall_pass", False))
    add_kv("Quality Score", f"{score}/100")
    add_kv("Grader Notes",  grader.get("grader_notes", ""))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    _font(p.add_run(
        "— This memo was drafted by the AI engine and reviewed by a qualified "
        "investment professional before IC presentation. —"
    ), size=8, italic=True, color=GREY)

    doc.save(out_path)
    return out_path


# ── Deal persistence ──────────────────────────────────────────────────────────
def _save_completed_deal(deal_id: str, deal: dict):
    try:
        existing = {}
        if COMPLETED_DEALS_FILE.exists():
            with open(COMPLETED_DEALS_FILE, "r", encoding="utf-8") as f:
                existing = json.load(f)
        existing[deal_id] = {k: v for k, v in deal.items()
                             if isinstance(v, (str, int, float, bool, dict, list, type(None)))}
        with open(COMPLETED_DEALS_FILE, "w", encoding="utf-8") as f:
            json.dump(existing, f, indent=2, default=str, ensure_ascii=False)
        print(f"[Pipeline] Saved {deal_id} to completed_deals.json")
    except Exception as e:
        print(f"[Pipeline] Could not save deal: {e}")


def _load_completed_deals() -> dict:
    if not COMPLETED_DEALS_FILE.exists():
        return {}
    try:
        with open(COMPLETED_DEALS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"[Startup] Could not load completed_deals.json: {e}")
        return {}


# ── In-memory stores ──────────────────────────────────────────────────────────
DEALS: dict = {}
AGENT_STATUS: dict = {
    "orchestrator":        {"status": "idle", "current_deal": None, "completed_at": None},
    "deal_sourcing":       {"status": "idle", "current_deal": None, "completed_at": None},
    "dd_research":         {"status": "idle", "current_deal": None, "completed_at": None},
    "financial_modelling": {"status": "idle", "current_deal": None, "completed_at": None},
    "ic_memo_drafting":    {"status": "idle", "current_deal": None, "completed_at": None},
    "outcomes_grader":     {"status": "idle", "current_deal": None, "completed_at": None},
}
ACTIVITY_LOG: list = [
    {"type": "grading",   "text": "IC memo ready — Greenway Highways passed QA (84/100). Partner review required.", "icon": "ti-check",     "bg": "var(--green-bg)",  "color": "var(--green)",  "agent": "Outcomes grader",           "ts": time.time() - 120},
    {"type": "memo",      "text": "Memo drafted — SunGrid Renewables 8-section memo complete. ESG 81/100.",         "icon": "ti-file-text", "bg": "var(--purple-bg)", "color": "var(--purple)", "agent": "IC memo agent",             "ts": time.time() - 1080},
    {"type": "modelling", "text": "Model complete — SunGrid base IRR 14.8% vs 14% IJF hurdle.",                     "icon": "ti-chart-bar", "bg": "#dbeafe",          "color": "#1e40af",       "agent": "Financial modelling agent", "ts": time.time() - 2040},
]

# Seed with demo deals
DEALS.update(DEMO_DEALS)

# Restore any completed pipeline deals from previous sessions
_completed = _load_completed_deals()
for _did, _deal in _completed.items():
    if _deal.get("docx_path") and not Path(_deal["docx_path"]).exists():
        _deal["docx_ready"] = False
    DEALS[_did] = _deal
    print(f"[Startup] Restored {_did} ({_deal.get('name','?')}) — score {_deal.get('quality_score','?')}/100")

# ── FastAPI app ───────────────────────────────────────────────────────────────
app = FastAPI(title="NIIFL AI Engine", version="1.0.0")
app.add_middleware(CORSMiddleware,
    allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
app.mount("/static", StaticFiles(directory=str(WEB_DIR / "static")), name="static")


# ── Request models ────────────────────────────────────────────────────────────
class DealSubmission(BaseModel):
    name: str
    fund: str
    sector: str
    ticket_cr: float
    stake_pct: float
    entry_ev_cr: float
    deal_lead: str
    brief: str
    revenue_fy22: Optional[float] = 0
    revenue_fy23: Optional[float] = 0
    revenue_fy24: Optional[float] = 0
    ebitda_fy22:  Optional[float] = 0
    ebitda_fy23:  Optional[float] = 0
    ebitda_fy24:  Optional[float] = 0
    net_debt_fy24:  Optional[float] = 0
    net_worth_fy24: Optional[float] = 0

class ActivityEvent(BaseModel):
    type: str
    text: str
    icon:  Optional[str] = "ti-check"
    bg:    Optional[str] = "var(--green-bg)"
    color: Optional[str] = "var(--green)"
    agent: Optional[str] = ""


# ── Demo pipeline runner ──────────────────────────────────────────────────────
async def run_pipeline_background(deal_id: str, submission: DealSubmission):
    """Runs the demo pipeline (demo_agent.py) in the background."""
    deal = DEALS[deal_id]
    hurdle_irr = _HURDLE.get(submission.fund, 12.0)

    for agent in AGENT_STATUS:
        AGENT_STATUS[agent]["status"] = "idle"
        AGENT_STATUS[agent]["completed_at"] = None
    AGENT_STATUS["orchestrator"]["status"] = "running"
    AGENT_STATUS["orchestrator"]["current_deal"] = deal["name"]
    AGENT_STATUS["deal_sourcing"]["status"] = "running"
    AGENT_STATUS["deal_sourcing"]["current_deal"] = deal["name"]
    AGENT_STATUS["deal_sourcing"]["completed_at"] = None
    deal["status"] = "screening"

    async def update_progress(stage: str, pct: int, msg: str = ""):
        _stage_map = {
            "screening": "deal_sourcing",
            "modelling": "financial_modelling",
            "memo":      "ic_memo_drafting",
            "grading":   "outcomes_grader",
        }
        agent_key = _stage_map.get(stage, stage)

        if "pipeline_stages" not in deal:
            deal["pipeline_stages"] = {}
        ps = deal["pipeline_stages"]
        if stage not in ps:
            ps[stage] = {"status": "pending", "started_at": None,
                         "completed_at": None, "duration_s": None, "msg": ""}
        ps[stage]["msg"] = msg

        if pct < 100:
            if ps[stage]["status"] != "running":
                ps[stage]["status"] = "running"
                ps[stage]["started_at"] = time.time()
            AGENT_STATUS[agent_key]["status"] = "running"
            AGENT_STATUS[agent_key]["current_deal"] = deal["name"]
            AGENT_STATUS[agent_key]["completed_at"] = None
            deal["status"] = {
                "screening": "screening",
                "modelling": "dd",
                "memo":      "ic_prep",
                "grading":   "ic_prep",
            }.get(stage, deal["status"])
        else:
            _now = time.time()
            ps[stage]["status"] = "done"
            ps[stage]["completed_at"] = _now
            if ps[stage].get("started_at"):
                ps[stage]["duration_s"] = round(_now - ps[stage]["started_at"], 1)
            AGENT_STATUS[agent_key]["status"] = "idle"
            AGENT_STATUS[agent_key]["current_deal"] = None
            AGENT_STATUS[agent_key]["completed_at"] = _now
            _ACT = {
                "screening": ("ti-search",    "var(--blue-bg)",   "var(--blue)",   "Deal sourcing agent"),
                "modelling": ("ti-chart-bar", "#dbeafe",          "#1e40af",       "Financial modelling agent"),
                "memo":      ("ti-file-text", "var(--purple-bg)", "var(--purple)", "IC memo agent"),
                "grading":   ("ti-check",     "var(--green-bg)",  "var(--green)",  "Outcomes grader"),
            }
            if stage in _ACT:
                icon, bg, color, agent_label = _ACT[stage]
                if stage == "screening":
                    act_text = f"Deal screened — {deal['name']} passed initial screening"
                elif stage == "modelling":
                    irr_str = msg.split("Base IRR ")[-1] if "Base IRR" in msg else "?"
                    act_text = f"Model complete — {deal['name']} base IRR {irr_str} vs {hurdle_irr}% hurdle"
                elif stage == "memo":
                    act_text = f"Memo drafted — {deal['name']} 8-section memo complete"
                else:
                    score_str = msg.split("score ")[-1] if "score" in msg else "?"
                    act_text = f"IC memo ready — {deal['name']} passed QA ({score_str}). Partner review required."
                ACTIVITY_LOG.append({
                    "type": stage, "text": act_text,
                    "icon": icon, "bg": bg, "color": color,
                    "agent": agent_label, "ts": _now,
                })

    try:
        from demo_agent import run_demo_pipeline
        financials = {
            "revenue_cr": {"FY22": submission.revenue_fy22, "FY23": submission.revenue_fy23, "FY24": submission.revenue_fy24},
            "ebitda_cr":  {"FY22": submission.ebitda_fy22,  "FY23": submission.ebitda_fy23,  "FY24": submission.ebitda_fy24},
        }
        result = await run_demo_pipeline(
            deal_name=submission.name,
            fund=submission.fund,
            sector=submission.sector,
            ticket_cr=submission.ticket_cr,
            stake_pct=submission.stake_pct,
            entry_ev_cr=submission.entry_ev_cr,
            brief=submission.brief,
            financials=financials,
            progress_callback=update_progress,
        )

        deal["memo_sections"] = result.get("memo_sections", {})
        deal["grader"]        = result.get("grader")
        deal["quality_score"] = result.get("quality_score")
        deal["irr_base"]      = result.get("irr_base")
        deal["irr_bull"]      = result.get("irr_bull")
        deal["irr_bear"]      = result.get("irr_bear")
        deal["irr_net"]       = result.get("irr_net")
        deal["moic"]          = result.get("moic")
        deal["memo_ready"]    = bool(deal.get("memo_sections"))
        deal["status"]        = "partner_review" if (result.get("grader") or {}).get("overall_pass") else "ic_prep"

        try:
            filename = _DEAL_DOCX_NAMES.get(deal_id, f"{deal_id}_IC_Memo.docx")
            out_path = str(OUTPUTS_DIR / filename)
            _build_docx(deal_id, deal, out_path)
            deal["docx_path"]  = out_path
            deal["docx_ready"] = True
            print(f"[Pipeline] DOCX written: {filename}")
        except Exception as e:
            deal["docx_error"] = str(e)
            print(f"[Pipeline] DOCX generation failed: {e}")

        _save_completed_deal(deal_id, deal)

    except Exception as e:
        deal["status"] = "error"
        deal["error"]  = str(e)
        print(f"[Pipeline] ERROR for '{submission.name}': {e}")
    finally:
        AGENT_STATUS["orchestrator"]["status"] = "idle"
        AGENT_STATUS["orchestrator"]["current_deal"] = None
        AGENT_STATUS["orchestrator"]["completed_at"] = time.time()


# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/")
async def serve_ui():
    return FileResponse(str(WEB_DIR / "static" / "index.html"))

@app.get("/api/deals")
async def list_deals():
    return JSONResponse([
        {k: v for k, v in d.items() if k not in ("memo_sections", "dd_result", "model_result")}
        for d in DEALS.values()
    ])

@app.post("/api/deals")
async def submit_deal(submission: DealSubmission, background_tasks: BackgroundTasks):
    deal_id = f"DC{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
    DEALS[deal_id] = {
        "deal_id": deal_id,
        "name": submission.name,
        "fund": submission.fund,
        "sector": submission.sector,
        "ticket_cr": submission.ticket_cr,
        "stake_pct": submission.stake_pct,
        "entry_ev_cr": submission.entry_ev_cr,
        "deal_lead": submission.deal_lead,
        "brief": submission.brief,
        "status": "queued",
        "quality_score": None,
        "irr_base": None,
        "moic": None,
        "memo_ready": False,
        "memo_sections": {},
        "grader": None,
        "submitted_at": datetime.datetime.now().isoformat(),
        "pipeline_stages": {
            "screening": {"status": "pending", "started_at": None, "completed_at": None, "duration_s": None, "msg": ""},
            "modelling": {"status": "pending", "started_at": None, "completed_at": None, "duration_s": None, "msg": ""},
            "memo":      {"status": "pending", "started_at": None, "completed_at": None, "duration_s": None, "msg": ""},
            "grading":   {"status": "pending", "started_at": None, "completed_at": None, "duration_s": None, "msg": ""},
        },
    }
    background_tasks.add_task(run_pipeline_background, deal_id, submission)
    return JSONResponse({"deal_id": deal_id, "status": "pipeline_started"})

@app.post("/api/submit-deal")
async def submit_deal_alias(submission: DealSubmission, background_tasks: BackgroundTasks):
    deal_id = f"DC{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
    DEALS[deal_id] = {
        "deal_id": deal_id,
        "name": submission.name,
        "fund": submission.fund,
        "sector": submission.sector,
        "ticket_cr": submission.ticket_cr,
        "stake_pct": submission.stake_pct,
        "entry_ev_cr": submission.entry_ev_cr,
        "deal_lead": submission.deal_lead,
        "brief": submission.brief,
        "status": "queued",
        "quality_score": None,
        "irr_base": None,
        "moic": None,
        "memo_ready": False,
        "memo_sections": {},
        "grader": None,
        "submitted_at": datetime.datetime.now().isoformat(),
        "pipeline_stages": {
            "screening": {"status": "pending", "started_at": None, "completed_at": None, "duration_s": None, "msg": ""},
            "modelling": {"status": "pending", "started_at": None, "completed_at": None, "duration_s": None, "msg": ""},
            "memo":      {"status": "pending", "started_at": None, "completed_at": None, "duration_s": None, "msg": ""},
            "grading":   {"status": "pending", "started_at": None, "completed_at": None, "duration_s": None, "msg": ""},
        },
    }
    background_tasks.add_task(run_pipeline_background, deal_id, submission)
    return JSONResponse({"deal_id": deal_id, "status": "pipeline_started"})

@app.get("/api/deals/{deal_id}")
async def get_deal(deal_id: str):
    if deal_id not in DEALS:
        raise HTTPException(404, "Deal not found")
    d = DEALS[deal_id]
    return JSONResponse({k: v for k, v in d.items() if k not in ("dd_result", "model_result")})

@app.get("/api/deals/{deal_id}/memo")
async def get_memo(deal_id: str):
    if deal_id not in DEALS:
        raise HTTPException(404, "Deal not found")
    d = DEALS[deal_id]
    return JSONResponse({
        "deal_id": deal_id,
        "name": d["name"],
        "fund": d["fund"],
        "memo_ready": d.get("memo_ready", False),
        "quality_score": d.get("quality_score"),
        "memo_sections": d.get("memo_sections", {}),
        "grader": d.get("grader", {}),
        "irr_base": d.get("irr_base"),
        "moic": d.get("moic"),
    })

@app.get("/api/deals/{deal_id}/memo/download")
async def download_memo(deal_id: str):
    if deal_id not in DEALS:
        raise HTTPException(404, "Deal not found")
    d = DEALS[deal_id]
    docx_path = d.get("docx_path")
    if not docx_path or not Path(docx_path).exists():
        filename  = _DEAL_DOCX_NAMES.get(deal_id, f"{deal_id}_IC_Memo.docx")
        candidate = OUTPUTS_DIR / filename
        if candidate.exists():
            docx_path = str(candidate)
        else:
            raise HTTPException(404, "DOCX not yet generated — wait for pipeline to complete")
    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"IC_Memo_{d['name'].replace(' ','_')}.docx",
    )

@app.get("/api/agents")
async def get_agents():
    return JSONResponse(AGENT_STATUS)

@app.get("/api/activity")
async def get_activity():
    return JSONResponse(list(reversed(ACTIVITY_LOG[-10:])))

@app.post("/api/activity")
async def add_activity(event: ActivityEvent):
    ACTIVITY_LOG.append({**event.dict(), "ts": time.time()})
    return JSONResponse({"ok": True})

@app.get("/api/system")
async def system_health():
    api_key   = os.getenv("ANTHROPIC_API_KEY", "")
    skills_ok = (ROOT_DIR / "skills" / "skill_fund_mandate.txt").exists()
    mcp_ok    = (ROOT_DIR / "mcp_servers" / "mcp_dealcloud.py").exists()
    outputs   = list(OUTPUTS_DIR.glob("*.json")) if OUTPUTS_DIR.exists() else []
    return JSONResponse({
        "api_key_set":     bool(api_key) and api_key != "YOUR_API_KEY_HERE",
        "api_key_preview": f"{api_key[:12]}...{api_key[-4:]}" if len(api_key) > 16 else "not set",
        "skills_ready":    skills_ok,
        "mcp_ready":       mcp_ok,
        "zdr_note":        "Enable ZDR at console.anthropic.com/settings",
        "data_residency":  "Local — configure Azure Central India for production",
        "deals_in_memory": len(DEALS),
        "output_files":    len(outputs),
        "model_orchestrator": "claude-opus-4-6",
        "model_subagent":     "claude-sonnet-4-6",
    })

@app.get("/api/outputs")
async def list_outputs():
    if not OUTPUTS_DIR.exists():
        return JSONResponse([])
    files = []
    for f in sorted(OUTPUTS_DIR.iterdir()):
        files.append({
            "name":     f.name,
            "size_kb":  round(f.stat().st_size / 1024, 1),
            "modified": datetime.datetime.fromtimestamp(f.stat().st_mtime).isoformat(),
        })
    return JSONResponse(files)

# ── Start server ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    print(f"\n{'='*50}")
    print(f"  AI Engine — Web Server")
    print(f"  Open: http://localhost:{port}")
    print(f"  API docs: http://localhost:{port}/docs")
    print(f"{'='*50}\n")
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True, app_dir=str(WEB_DIR))
